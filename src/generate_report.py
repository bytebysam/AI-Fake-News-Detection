from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUTPUT_PATH = Path("docs/AI_Fake_News_Detection_Report.docx")

document = Document()

section = document.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = document.styles

styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(10)

for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Times New Roman"


def add_heading(text, level=1):
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph(text, bold_prefix=None):
    paragraph = document.add_paragraph()

    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

title_run = title.add_run(
    "AI-Powered Fake News Detection Using Text Classification"
)
title_run.bold = True
title_run.font.name = "Times New Roman"
title_run.font.size = Pt(18)

author = document.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author.add_run("bytebysam\n")
author.add_run("Computer Science and Engineering")

add_heading("Abstract", 1)

add_paragraph(
    "The rapid growth of digital media and online news platforms has "
    "increased the spread of misinformation and fake news. Manual "
    "verification of large volumes of news articles is time-consuming "
    "and difficult. This project presents a machine learning-based fake "
    "news detection system using Natural Language Processing and text "
    "classification techniques. The WELFake dataset was used for model "
    "development. Text preprocessing, TF-IDF feature extraction, and "
    "four classification algorithms—K-Nearest Neighbors, Logistic "
    "Regression, Random Forest, and a Neural Network—were implemented "
    "and evaluated. Experimental results showed that Random Forest "
    "achieved the highest accuracy of 95.58%. The developed system "
    "demonstrates the effectiveness of machine learning and NLP "
    "techniques for automatic fake news classification."
)

keywords = document.add_paragraph()
keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

run = keywords.add_run("Keywords—")
run.bold = True

keywords.add_run(
    "Fake News Detection, Natural Language Processing, Machine Learning, "
    "TF-IDF, Random Forest, Text Classification."
)

add_heading("I. INTRODUCTION", 1)

add_paragraph(
    "Fake news refers to false or misleading information presented in "
    "the form of legitimate news. The widespread use of social media "
    "and digital news platforms has made the rapid distribution of "
    "misinformation a significant problem."
)

add_paragraph(
    "Manual fact-checking requires considerable time and human effort. "
    "Artificial Intelligence and Machine Learning can assist in "
    "automatically identifying linguistic patterns associated with fake "
    "and real news articles."
)

add_paragraph(
    "The objective of this project is to design and implement a machine "
    "learning pipeline capable of classifying news articles as fake or "
    "real using Natural Language Processing and text classification."
)

add_heading("II. DATASET DESCRIPTION", 1)

add_paragraph(
    "The WELFake dataset was used in this project. The original dataset "
    "contained 72,134 news articles with article title, textual content, "
    "and a binary classification label."
)

add_paragraph(
    "Label 0 represents Fake News and label 1 represents Real News. "
    "The dataset contained 35,028 fake news articles and 37,106 real "
    "news articles."
)

add_paragraph(
    "The original data contained 558 missing title values and 39 missing "
    "text values. After preprocessing and removal of empty cleaned "
    "articles, 72,080 samples remained for model development."
)

add_heading("III. METHODOLOGY", 1)

add_heading("A. Text Preprocessing", 2)

add_paragraph(
    "The title and article text were combined into a single textual "
    "feature. Text was converted to lowercase, HTML tags and URLs were "
    "removed, and punctuation and special characters were eliminated. "
    "English stopwords were removed and WordNet lemmatization was "
    "applied using the NLTK library."
)

add_heading("B. Feature Extraction", 2)

add_paragraph(
    "Term Frequency-Inverse Document Frequency (TF-IDF) was used to "
    "transform textual data into numerical feature vectors. The "
    "vectorizer was configured with a maximum of 5,000 features. The "
    "resulting feature matrix contained 72,080 samples and 5,000 "
    "features."
)

add_heading("C. Train-Test Split", 2)

add_paragraph(
    "The processed dataset was divided into training and testing sets "
    "using an 80:20 ratio with stratified sampling. The training set "
    "contained 57,664 samples and the testing set contained 14,416 "
    "samples."
)

add_heading("D. Classification Models", 2)

add_paragraph(
    "Four classification models were implemented: K-Nearest Neighbors "
    "(KNN), Logistic Regression, Random Forest, and a Multi-Layer "
    "Perceptron Neural Network."
)

add_paragraph(
    "The KNN classifier used K=5. Random Forest used 100 decision trees. "
    "The Neural Network contained one hidden layer with 100 neurons and "
    "used early stopping to reduce unnecessary training."
)

add_heading("IV. EVALUATION METRICS", 1)

add_paragraph(
    "The models were evaluated using accuracy, precision, recall, and "
    "F1-score. Confusion matrices were also generated to analyze the "
    "classification performance of individual models."
)

add_heading("V. RESULTS", 1)

table = document.add_table(rows=1, cols=5)
table.style = "Table Grid"

headers = [
    "Model",
    "Accuracy",
    "Precision",
    "Recall",
    "F1-Score",
]

for index, header in enumerate(headers):
    table.rows[0].cells[index].text = header

results = [
    ["Random Forest", "95.58%", "94.57%", "96.98%", "95.76%"],
    ["Neural Network", "94.76%", "94.55%", "95.30%", "94.93%"],
    ["Logistic Regression", "94.58%", "94.32%", "95.18%", "94.75%"],
    ["KNN", "71.37%", "65.29%", "94.56%", "77.25%"],
]

for result in results:
    cells = table.add_row().cells

    for index, value in enumerate(result):
        cells[index].text = value

caption = document.add_paragraph(
    "TABLE I. PERFORMANCE COMPARISON OF CLASSIFICATION MODELS"
)
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(
    "Random Forest achieved the highest classification accuracy of "
    "95.58% and the highest F1-score of 95.76% among the evaluated "
    "models."
)

comparison_image = Path("outputs/model_accuracy_comparison.png")

if comparison_image.exists():
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph.add_run().add_picture(
        str(comparison_image),
        width=Inches(5.8),
    )

    figure_caption = document.add_paragraph(
        "Fig. 1. Accuracy comparison of the evaluated models."
    )
    figure_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading("VI. DISCUSSION", 1)

add_paragraph(
    "The experimental results demonstrate that ensemble and linear "
    "classification techniques perform effectively on TF-IDF text "
    "features. Random Forest achieved the best overall performance. "
    "The combination of multiple decision trees enabled the model to "
    "identify complex patterns within the feature space."
)

add_paragraph(
    "Logistic Regression achieved strong performance despite its "
    "relatively simple architecture. The Neural Network achieved an "
    "accuracy of 94.76%, with early stopping terminating training after "
    "validation performance stopped improving."
)

add_paragraph(
    "KNN achieved the lowest accuracy of 71.37%. Its high recall and "
    "comparatively low precision indicate that distance-based "
    "classification is less suitable for the high-dimensional sparse "
    "TF-IDF feature space used in this project."
)

add_heading("VII. APPLICATION DEVELOPMENT", 1)

add_paragraph(
    "A Streamlit-based web application named TruthLens AI was developed "
    "to provide an interactive interface for the trained classification "
    "system. Users can enter news article text and receive a fake or "
    "real news pattern prediction together with the model confidence."
)

add_paragraph(
    "The application uses the trained Random Forest classifier and the "
    "saved TF-IDF vectorizer. The application performs linguistic "
    "pattern classification and does not perform real-time factual "
    "verification."
)

add_heading("VIII. LIMITATIONS", 1)

add_paragraph(
    "The system classifies articles based on linguistic patterns learned "
    "from the WELFake dataset. It does not verify claims against live "
    "news sources or external fact-checking databases. Predictions may "
    "also be less reliable for short headlines or text that differs "
    "significantly from the training distribution."
)

add_heading("IX. CONCLUSION", 1)

add_paragraph(
    "This project successfully developed an AI-powered fake news "
    "detection system using Natural Language Processing and Machine "
    "Learning. A complete text classification pipeline was implemented, "
    "including preprocessing, TF-IDF feature extraction, model training, "
    "evaluation, comparison, and interactive application development."
)

add_paragraph(
    "Among the four evaluated models, Random Forest achieved the highest "
    "accuracy of 95.58%. The results demonstrate that NLP and machine "
    "learning techniques can effectively identify linguistic patterns "
    "associated with fake and real news articles."
)

add_heading("X. FUTURE SCOPE", 1)

future_points = [
    "Implementation of transformer-based models such as BERT.",
    "Integration with real-time fact-checking services.",
    "Development of multilingual fake news detection.",
    "Use of Explainable AI techniques.",
    "Online deployment of the TruthLens AI application.",
]

for point in future_points:
    document.add_paragraph(point, style="List Bullet")

add_heading("REFERENCES", 1)

references = [
    "[1] WELFake Dataset, Fake News Classification Dataset.",
    "[2] Scikit-learn, Machine Learning in Python.",
    "[3] NLTK Project, Natural Language Toolkit.",
    "[4] Streamlit, Open-source Python Application Framework.",
]

for reference in references:
    document.add_paragraph(reference)

OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

document.save(OUTPUT_PATH)

print("Report generated successfully!")
print("Saved to:", OUTPUT_PATH)